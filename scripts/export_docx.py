"""
export_docx.py — （選用）把譯稿另外輸出成 Word .docx

讀跟 combine 相同的 build/sec*.html（+ build/meta.html），產出 .docx。
docx 是靜態文件、沒有三態切換，所以用 --view 決定輸出哪個版本：
  zh    純中文（預設，最適合閱讀/交稿）
  both  中英對照（每段先中文，接一段灰階原文）
  orig  純原文

相依：python-docx（pip install python-docx）
圖片/公式若是 base64 內嵌會一併寫入；掃描原頁的大圖也支援。

用法：
    python export_docx.py --build build --out "out/論文_中譯.docx" --view zh
    python export_docx.py --build build --out "out/論文_對照.docx" --view both
"""
import argparse
import base64
import re
import sys
from html import unescape
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    sys.exit("需要 python-docx：pip install python-docx")

sys.path.insert(0, str(Path(__file__).parent))
from paper_utils import extract_sections

GREY = RGBColor(0x7c, 0x76, 0x6c)
ACCENT = RGBColor(0x2f, 0x5b, 0x8a)


def strip_tags(html):
    """去掉標籤但保留文字；<span class="en">x</span> → （x）。"""
    html = re.sub(r'<span class="en">(.*?)</span>', r'（\1）', html, flags=re.DOTALL)
    html = re.sub(r'<br\s*/?>', '\n', html)
    html = re.sub(r'<[^>]+>', '', html)
    return unescape(html).strip()


def inner(html, cls):
    """取某 class 的 div 內層 HTML（第一個）。"""
    m = re.search(r'<div class="%s"[^>]*>(.*?)</div>' % cls, html, flags=re.DOTALL)
    return m.group(1) if m else None


def add_para(doc, text, *, size=11.5, color=None, italic=False, align=None,
             space_after=8, style_font="Noto Serif TC"):
    if text is None:
        return
    p = doc.add_paragraph()
    if align:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    for i, line in enumerate(text.split("\n")):
        if i:
            p.add_run().add_break()
        r = p.add_run(line)
        r.font.size = Pt(size)
        r.font.name = style_font
        if color:
            r.font.color.rgb = color
        r.font.italic = italic
    return p


def add_image_b64(doc, data_uri, width_in=5.6):
    try:
        b = base64.b64decode(data_uri.split(",", 1)[1])
        import io
        doc.add_picture(io.BytesIO(b), width=Inches(width_in))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    except Exception as e:
        print(f"  ⚠ 圖片寫入失敗：{e}")


# 依出現順序處理 section 內各區塊
BLOCK_RE = re.compile(
    r'<h1[^>]*>(?P<h1>.*?)</h1>'
    r'|<h2[^>]*>(?P<h2>.*?)</h2>'
    r'|<h3[^>]*>(?P<h3>.*?)</h3>'
    r'|<div class="pgmark"[^>]*>(?P<pg>.*?)</div>'
    r'|<div class="para"[^>]*>(?P<para>.*?)</div>\s*</div>'
    r'|<figure[^>]*>(?P<fig>.*?)</figure>'
    r'|<div class="poem"[^>]*>(?P<poem>.*?)</div>',
    re.DOTALL)


def render_section(doc, sec_html, view):
    for m in BLOCK_RE.finditer(sec_html):
        if m.group('h1') is not None:
            add_para(doc, strip_tags(m.group('h1')), size=18, color=ACCENT,
                     style_font="Noto Sans TC", space_after=10)
        elif m.group('h2') is not None:
            add_para(doc, strip_tags(m.group('h2')), size=14, color=ACCENT,
                     style_font="Noto Sans TC", space_after=6)
        elif m.group('h3') is not None:
            add_para(doc, strip_tags(m.group('h3')), size=12,
                     style_font="Noto Sans TC", space_after=4)
        elif m.group('pg') is not None:
            add_para(doc, "— " + strip_tags(m.group('pg')) + " —", size=9,
                     color=GREY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)
        elif m.group('para') is not None:
            # regex 把 orig 的閉合 </div> 吃掉了，補回再抽
            full = '<div class="para">' + m.group('para') + '</div></div>'
            zh = inner(full, 'zh')
            orig = inner(full, 'orig')
            if view in ('zh', 'both') and zh:
                add_para(doc, strip_tags(zh), size=11.5)
            if view in ('orig', 'both') and orig:
                add_para(doc, strip_tags(orig), size=10.5, color=GREY, italic=True,
                         space_after=12 if view == 'both' else 8)
        elif m.group('fig') is not None:
            figu = m.group('fig')
            src = re.search(r'src="(data:image[^"]+)"', figu)
            if src:
                add_image_b64(doc, src.group(1))
            cap = re.search(r'<figcaption[^>]*>(.*?)</figcaption>', figu, re.DOTALL)
            if cap:
                add_para(doc, strip_tags(cap.group(1)), size=9.5, color=GREY,
                         align=WD_ALIGN_PARAGRAPH.CENTER)
        elif m.group('poem') is not None:
            poem = m.group('poem')
            zh = inner(poem + '</div>', 'zh')
            orig = inner(poem + '</div>', 'orig')
            if view in ('zh', 'both') and zh:
                add_para(doc, strip_tags(zh), size=11, italic=True, space_after=4)
            if view in ('orig', 'both') and orig:
                add_para(doc, strip_tags(orig), size=10.5, color=GREY, italic=True)
            if not zh and not orig:
                add_para(doc, strip_tags(poem), size=10.5, color=GREY, italic=True)


def render_meta(doc, meta_html):
    t = re.search(r'<h1 class="dm-title"[^>]*>(.*?)</h1>', meta_html, re.DOTALL)
    if t:
        add_para(doc, strip_tags(t.group(1)), size=20, color=ACCENT,
                 style_font="Noto Sans TC", space_after=6)
    for fm in re.finditer(r'<dt>(.*?)</dt>\s*<dd>(.*?)</dd>', meta_html, re.DOTALL):
        add_para(doc, f"{strip_tags(fm.group(1))}：{strip_tags(fm.group(2))}",
                 size=10.5, color=GREY, space_after=2)
    note = re.search(r'<p class="dm-note"[^>]*>(.*?)</p>', meta_html, re.DOTALL)
    if note:
        add_para(doc, strip_tags(note.group(1)), size=9.5, color=GREY, space_after=14)
    doc.add_paragraph()


def main():
    ap = argparse.ArgumentParser(description="譯稿輸出為 Word .docx")
    ap.add_argument("--build", default="build")
    ap.add_argument("--out", required=True)
    ap.add_argument("--view", default="zh", choices=["zh", "both", "orig"],
                    help="輸出版本（docx 無切換，預設純中文）")
    ap.add_argument("--meta", help="書目 header（預設自動找 build/meta.html）")
    ap.add_argument("--no-credit", action="store_true")
    args = ap.parse_args()

    build_dir = Path(args.build)
    files = sorted(build_dir.glob("sec*.html"))
    if not files:
        sys.exit(f"找不到 {build_dir}/sec*.html")

    doc = Document()
    doc.styles['Normal'].font.name = "Noto Serif TC"
    doc.styles['Normal'].font.size = Pt(11.5)

    meta_path = Path(args.meta) if args.meta else (build_dir / "meta.html")
    if meta_path.exists():
        render_meta(doc, meta_path.read_text(encoding="utf-8"))
        print(f"書目 header：{meta_path.name} ✓")

    nsec = 0
    for f in files:
        for sec in extract_sections(f.read_text(encoding="utf-8")):
            render_section(doc, sec, args.view)
            nsec += 1

    if not args.no_credit:
        doc.add_paragraph()
        add_para(doc,
                 "本譯本以開源翻譯工具製作，原始碼見 github.com/Zaious/translate-academic-paper，歡迎自由取用、修改與散布。",
                 size=9, color=GREY, align=WD_ALIGN_PARAGRAPH.CENTER)

    out = Path(args.out)
    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out))
    print(f"合併 {len(files)} 檔、{nsec} 個 section，view={args.view}")
    print(f"Done → {out}  ({out.stat().st_size // 1024} KB)")


if __name__ == "__main__":
    main()
